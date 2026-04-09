import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def set_cell_background(cell, color):
    # color in hex without #
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, text_color=None):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = 'Arial' # match or fallback
            if bold:
                r.font.bold = True
            if text_color:
                r.font.color.rgb = RGBColor.from_string(text_color)

doc = docx.Document('AI_Pricing_Estimate.docx')
target_p = None

for p in doc.paragraphs:
    if '3A.' in p.text and 'Meta WhatsApp Fees' in p.text:
        target_p = p
        break

if not target_p:
    print("Could not find Target paragraph '3A.'")
else:
    print("Found target paragraph:", target_p.text)
    
    # insert a paragraph before target to separate
    new_p = doc.add_paragraph('Cost Comparison of Broadcasting Platforms')
    new_p_bold = new_p.runs[0] if new_p.runs else new_p.add_run('Cost Comparison of Broadcasting Platforms')
    new_p_bold.bold = True
    target_p._p.addprevious(new_p._p)

    # create table
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Normal Table'
    table.autofit = True
    
    headers = ['Platform', 'Setup/Monthly Fees', 'Message Markup (per msg)', 'Service Conv. Cost']
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, text_color='FFFFFF')
        set_cell_background(table.cell(0, i), '2E75B6')
        
    data = [
        ['Meta Cloud API', 'Free', '0x (Original Meta Rates)', 'Free'],
        ['Twilio', '~/mo for number', '+ .005 (Markup)', 'Billed at Twilio rates'],
        ['Gupshup', 'Subscription Required', 'Markup applied', 'Billed at Gupshup rates'],
    ]
    
    for r_idx, row_data in enumerate(data, start=1):
        for c_idx, text in enumerate(row_data):
            set_cell_text(table.cell(r_idx, c_idx), text)
            
    # Move table before target_p
    target_p._p.addprevious(table._tbl)
    
    empty_p = doc.add_paragraph('')
    target_p._p.addprevious(empty_p._p)
    
    doc.save('AI_Pricing_Estimate_Modified.docx')
    print("Table inserted and saved to AI_Pricing_Estimate_Modified.docx")
