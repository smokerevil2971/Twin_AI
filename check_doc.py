import docx

doc = docx.Document('AI_Pricing_Estimate.docx')
for i, p in enumerate(doc.paragraphs):
    if 'WhatsApp' in p.text or 'Cost' in p.text or 'Messaging' in p.text:
        print(f"Paragraph {i}: {p.text.encode('ascii', 'ignore').decode()}")

print("\n--- TABLES ---")
for i, t in enumerate(doc.tables):
    print(f"Table {i} style: {t.style.name}")
    try:
        # Check first row cell colors if any exist
        cell = t.cell(0, 0)
        shading_elm = cell._tc.get_or_add_tcPr().find('.//w:shd', namespaces=cell._tc.nsmap)
        bg = shading_elm.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') if shading_elm is not None else 'None'
        print(f"  Bg check: {bg}")
    except Exception as e:
        print(f"  Bg error: {e}")
