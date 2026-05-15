"""
Knowledge Base ingestion service — Phase 3.1

Pipeline:
  file bytes → extract text → chunk → embed (NIM or Gemini) → store (pgvector + Postgres)
"""

import asyncio
import uuid
from datetime import datetime, timezone
from typing import Optional, Any

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
import os



from langchain.text_splitter import RecursiveCharacterTextSplitter
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, func, delete as sa_delete
from fastapi import HTTPException

from core.config import settings
from models.models import KnowledgeBase, KnowledgeChunk

from core.logging import logger

# ─── Constants ────────────────────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".txt", ".md", ".docx"}
MAX_FILE_SIZE_BYTES = settings.max_upload_size_mb * 1024 * 1024



# ─── Text extraction ──────────────────────────────────────────────────────────


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract raw text from PDF, image, or plain text file.

    For PDFs: first tries native text extraction (fast, works for digital PDFs).
    If the PDF is scanned/image-based and returns no text, falls back to
    rendering each page as an image and running pytesseract OCR on it.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext == ".pdf":
        text_parts = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text())
        text = "\n".join(text_parts)

        # Fallback: scanned / image-based PDF — render each page and OCR it
        if not text.strip():
            logger.info(
                f"[KB] '{filename}' has no selectable text — attempting OCR on pages"
            )
            ocr_parts = []
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page_num, page in enumerate(doc):
                    # Render page at 2x resolution for better OCR accuracy
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_bytes = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_bytes))
                    page_text = pytesseract.image_to_string(image)
                    if page_text.strip():
                        ocr_parts.append(page_text)
                    logger.info(f"[KB] OCR page {page_num + 1}: {len(page_text)} chars")
            text = "\n".join(ocr_parts)

    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image)

    elif ext in {".txt", ".md"}:
        text = file_bytes.decode("utf-8", errors="replace")

    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        logger.info(f"[KB] Extracted {len(parts)} paragraphs/rows from Word document")

    else:
        raise HTTPException(400, f"Unsupported file type: {ext}")

    if not text.strip():
        raise HTTPException(
            422,
            "No text could be extracted from the file. "
            "If this is a scanned PDF, ensure the scan quality is clear. "
            "You can also try uploading it as individual page images (.jpg/.png).",
        )

    logger.info(f"[KB] Extracted {len(text)} chars from {filename}")
    return text


# ─── Semantic chunking ────────────────────────────────────────────────────────


def chunk_text(text: str) -> list[str]:
    """Split text into chunks — size and overlap configurable via KB_CHUNK_SIZE/KB_CHUNK_OVERLAP."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.kb_chunk_size,
        chunk_overlap=settings.kb_chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_text(text)
    logger.info(f"[KB] Split into {len(chunks)} chunks")
    return chunks


# ─── Postgres/pgvector connection ──────────────────────────────────────────────


async def query_knowledge_base_pgvector(
    db: AsyncSession,
    query_embedding: list[float],
    n_results: int = 15,
) -> dict:
    """
    Query PostgreSQL (pgvector) for top-k chunks most similar to the query embedding.
    Only returns chunks where the parent knowledge base is_active == True.
    """
    try:
        distance = KnowledgeChunk.embedding.cosine_distance(query_embedding).label("distance")
        q = (
            select(KnowledgeChunk, distance)
            .join(KnowledgeBase, KnowledgeChunk.knowledge_base_id == KnowledgeBase.id)
            .where(KnowledgeBase.is_active == True)
            .order_by(distance)
            .limit(n_results)
        )
        
        results = (await db.execute(q)).all()
        
        docs = []
        dists = []
        for chunk, dist in results:
            docs.append(chunk.content)
            dists.append(float(dist))
            
        return {"documents": docs, "distances": dists}
    except Exception as e:
        logger.error(f"[KB] pgvector query failed: {e}")
        return {"documents": [], "distances": []}


# ─── Embeddings (dual-provider) ───────────────────────────────────────────────


def embed_texts(texts: list[str]) -> list[list[float]]:
    """
    Generate embeddings using the configured provider.
    LLM_PROVIDER=nim  → NVIDIA NIM nemoretriever (OpenAI-compatible)
    LLM_PROVIDER=gemini → Google Gemini embedding-001
    """
    if settings.is_nim:
        return _embed_texts_nim(texts)
    else:
        return _embed_texts_gemini(texts)


def _embed_texts_nim(texts: list[str]) -> list[list[float]]:
    """
    Embed via NVIDIA NIM using direct HTTP (bypasses OpenAI SDK quirks).
    Uses the LLM API key — NIM keys are universal across all models.
    """
    import httpx, json

    url = f"{settings.nim_base_url}/embeddings"
    headers = {
        "Authorization": f"Bearer {settings.nim_embed_api_key or settings.nim_llm_api_key}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    embeddings = []
    batch_size = 50
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        payload = {
            # MED-05 fix: Was hardcoded to "nvidia/llama-3.2-nemoretriever-300m-embed-v1".
            # If EMBEDDING_MODEL is changed in .env for the RAG query path but not here,
            # ingestion uses a different-dimension model → ChromaDB rejects vectors at
            # query time with "dimension mismatch" errors.
            "model": settings.embedding_model,
            "input": batch,
            "input_type": "passage",
            "encoding_format": "float",
            "truncate": "END",
        }
        with httpx.Client(timeout=30) as client:
            resp = client.post(url, headers=headers, json=payload)
            if not resp.is_success:
                logger.error(
                    f"[KB][NIM] Embed error {resp.status_code}: {resp.text[:200]}"
                )
                resp.raise_for_status()
            data = resp.json()
        embeddings.extend([item["embedding"] for item in data["data"]])
    logger.info(f"[KB][NIM] Generated {len(embeddings)} embeddings")
    return embeddings


def _embed_texts_gemini(texts: list[str]) -> list[list[float]]:
    """Embed via Google Gemini embedding-001."""
    import google.generativeai as genai


    embeddings = []
    batch_size = 100
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        result = genai.embed_content(
            model=settings.embedding_model,
            content=batch,
            task_type="retrieval_document",
        )
        embeddings.extend(result["embedding"])
    logger.info(
        f"[KB][Gemini] Generated {len(embeddings)} embeddings via {settings.embedding_model}"
    )
    return embeddings


# ─── Core ingestion ───────────────────────────────────────────────────────────


async def ingest_document(
    db: AsyncSession,
    file_bytes: bytes,
    filename: str,
    category: str,
    valid_from: Optional[datetime] = None,
    valid_until: Optional[datetime] = None,
) -> dict:
    """
    Full ingestion pipeline:
    1. Extract text
    2. Chunk text
    3. Generate embeddings (NIM or Gemini)
    4. Store chunks in pgvector
    5. Save KnowledgeBase record in Postgres
    """
    # Validate file size
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(413, f"File exceeds {settings.max_upload_size_mb}MB limit.")

    # Validate / default offers dates
    if category == "offers":
        from datetime import date

        if not valid_from:
            valid_from = datetime.now(timezone.utc).replace(
                hour=0, minute=0, second=0, microsecond=0
            )
        if not valid_until:
            today = date.today()
            valid_until = datetime(today.year, 12, 31, 23, 59, 59, tzinfo=timezone.utc)

    # Extract text — CPU-bound (PyMuPDF + pytesseract OCR).
    # Run in thread pool to avoid blocking the async event loop.
    try:
        text = await asyncio.to_thread(extract_text, file_bytes, filename)
    except HTTPException:
        raise  # re-raise clean 400/422 from extract_text
    except Exception as e:
        logger.error(f"[KB] Text extraction failed: {e}")
        raise HTTPException(422, f"Could not extract text from file: {str(e)}")

    # Chunk (fast, in-process — no threading needed)
    chunks = chunk_text(text)
    if not chunks:
        raise HTTPException(422, "File produced no usable text chunks.")

    # Embed — synchronous HTTP calls to NIM/Gemini.
    # Run in thread pool to avoid blocking the async event loop.
    try:
        embeddings = await asyncio.to_thread(embed_texts, chunks)
    except Exception as e:
        logger.error(f"[KB] Embedding failed: {e}")
        raise HTTPException(502, f"Embedding service error: {str(e)}")

    doc_id = uuid.uuid4()

    # ── Layer 5: Privacy Guardrail — redact PII before storing ────────────────
    try:
        from services.guardrails.privacy_guard import scan_and_redact
        cleaned_chunks = [scan_and_redact(chunk) for chunk in chunks]
        pii_count = sum(1 for orig, clean in zip(chunks, cleaned_chunks) if orig != clean)
        if pii_count:
            logger.warning(
                f"[GUARDRAIL][PRIVACY] Redacted PII in {pii_count}/{len(chunks)} "
                f"chunks of '{filename}' before pgvector ingestion"
            )
        chunks = cleaned_chunks
    except Exception as _pe:
        logger.warning(f"[GUARDRAIL][PRIVACY] PII scan failed (non-fatal): {_pe}")

    metadata_list = [
        {
            "doc_id": str(doc_id),
            "category": category,
            "filename": filename,
            "valid_from": valid_from.isoformat() if valid_from else "",
            "valid_until": valid_until.isoformat() if valid_until else "",
            "is_active": "true",
            "chunk_index": i,
        }
        for i in range(len(chunks))
    ]

    # Store chunks in pgvector
    import json
    pg_chunks = [
        KnowledgeChunk(
            knowledge_base_id=doc_id,
            content=chunk,
            embedding=embeddings[i],
            chunk_metadata=json.dumps(metadata_list[i])
        )
        for i, chunk in enumerate(chunks)
    ]
    db.add_all(pg_chunks)
    logger.info(f"[KB] Indexed {len(chunks)} chunks into pgvector")

    # Save Postgres record
    kb_record = KnowledgeBase(
        id=doc_id,
        filename=filename,
        category=category,
        chroma_ids=[],
        valid_from=valid_from,
        valid_until=valid_until,
        is_active=True,
    )
    db.add(kb_record)
    

    await db.commit()
    await db.refresh(kb_record)

    return {
        "id": str(kb_record.id),
        "filename": kb_record.filename,
        "category": kb_record.category,
        "chunks_indexed": len(chunks),
        "valid_from": kb_record.valid_from.isoformat()
        if kb_record.valid_from
        else None,
        "valid_until": kb_record.valid_until.isoformat()
        if kb_record.valid_until
        else None,
        "is_active": kb_record.is_active,
        "created_at": kb_record.created_at.isoformat(),
        "status": "indexed",
    }


# ─── List documents ───────────────────────────────────────────────────────────


async def list_documents(
    db: AsyncSession,
    category: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
) -> dict:
    q = select(KnowledgeBase)
    if category:
        q = q.where(KnowledgeBase.category == category)

    total = (
        await db.execute(select(func.count()).select_from(q.subquery()))
    ).scalar_one()
    rows = (
        (
            await db.execute(
                q.order_by(KnowledgeBase.created_at.desc())
                .offset((page - 1) * page_size)
                .limit(page_size)
            )
        )
        .scalars()
        .all()
    )

    return {
        "documents": [_kb_dict(r) for r in rows],
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": max(1, (total + page_size - 1) // page_size),
    }


# ─── Delete document ──────────────────────────────────────────────────────────


async def delete_document(
    db: AsyncSession,
    doc_id: uuid.UUID,
) -> dict:
    result = await db.execute(
        select(KnowledgeBase).where(
            KnowledgeBase.id == doc_id,
        )
    )
    record = result.scalar_one_or_none()
    if not record:
        raise HTTPException(404, "Document not found.")

    # Delete pgvector chunks only
    logger.info(f"[KB] Deleting pgvector chunks for doc {doc_id}")

    # Delete pgvector chunks
    await db.execute(
        sa_delete(KnowledgeChunk).where(KnowledgeChunk.knowledge_base_id == doc_id)
    )

    # Mark inactive in Postgres
    await db.execute(
        update(KnowledgeBase).where(KnowledgeBase.id == doc_id).values(is_active=False)
    )
    await db.commit()

    return {"id": str(doc_id), "status": "deleted"}


# ─── Clear all documents ──────────────────────────────────────────────────────


async def clear_all_documents(db: AsyncSession) -> dict:
    """
    Wipe the entire knowledge base:
    1. Hard-delete all KnowledgeChunk rows from pgvector.
    2. Hard-delete all KnowledgeBase rows from Postgres.

    Returns a summary dict with counts.
    """
    # Count rows before deletion for the response
    total = (await db.execute(select(func.count()).select_from(KnowledgeBase))).scalar_one()

    # ── pgvector + Postgres: hard-delete all records ────────────────────────────
    await db.execute(sa_delete(KnowledgeChunk))
    logger.info(f"[KB] Deleted all KnowledgeChunk rows from pgvector")

    # ── Postgres: hard-delete all KB records ──────────────────────────────────
    await db.execute(sa_delete(KnowledgeBase))
    await db.commit()
    logger.info(f"[KB] Deleted {total} KnowledgeBase records from Postgres")

    return {
        "status": "cleared",
        "documents_removed": total,
        "message": f"Knowledge base cleared. {total} document(s) removed.",
    }


# ─── Helper ───────────────────────────────────────────────────────────────────


def _kb_dict(r: KnowledgeBase) -> dict:
    return {
        "id": str(r.id),
        "filename": r.filename,
        "category": r.category,
        "chunks_count": len(r.chroma_ids) if r.chroma_ids else 0,
        "valid_from": r.valid_from.isoformat() if r.valid_from else None,
        "valid_until": r.valid_until.isoformat() if r.valid_until else None,
        "is_active": r.is_active,
        "created_at": r.created_at.isoformat(),
    }
